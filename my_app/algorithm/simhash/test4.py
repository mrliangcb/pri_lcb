import re
import docx
import collections
from collections import namedtuple as nt
# -*- coding: utf-8 -*-
# coding:utf-8
# coding=gbk
para_obj =nt('paragraph', ['type', 'position', 'origin','str_','flag','test','para_num','from_global']) # flag和test怎么用
para_obj.__new__.__defaults__ = ('para',None, None,None,None,None,None)
x=['123','45678','91011121314']

#一个元素，至少是空句子，其次是空段



path1=r'D:\lcb_note\code\Program\10月项目\查重需求资料\查重需求资料\1.招标文件文档查重对比材料\招标文件-基于NLP的商务文本数据清洗关键技术研究（1021评审）（第一版）.docx'
# path1=r'D:\lcb_note\code\Program\10月项目\my_docx\基于NLP的商务文本数据清洗关键技术研究项目合同+-+-打印版.docx'
doc1=docx.Document(path1)

base = r"D:\lcb_note\code\Program\10月项目\查重需求资料\查重需求资料\1.招标文件文档查重对比材料\招标文件 CWEME-1911ZSWZ-2J039 基于NLP的商务文本数据清洗关键技术研究项目-2019年12月中国水利电力物资集团有限公司项目（第三版终版）.docx"
# path2 = r'D:\lcb_note\code\Program\10月项目\my_docx\招标文件 CWEME-1911ZSWZ-2J039 基于NLP的商务文本数据清洗关键技术研究项目-2019年12月中国水利电力物资集团有限公司项目（第三版终版）.docx'
doc2 = docx.Document(base)



def exctract_heading(para_list):
    heading_list=[]
    para_num = -1
    global_obj=[]
    pos_num=0
    for i,para in enumerate(para_list):
        text=para.text.strip()
        if (text not in trasbin) and text:
            # 有两个解析，等判空之后再做，减少一点时间，剪枝
            type_name = para.style.name
            str_split = text.split(' ')[-1]
            is_heading = 0
            #校正章号
            if type_name.startswith('Heading'):
                # if (not para.style.name.startswith('Normal')) and (not para.style.name.startswith('normal')):
                ptr = r'第(.*?)章'  # 非贪心
                result = re.findall(ptr, text)
                if result and result[0] != '':
                    para_num += 1
                    # 这是一个章标题
                    # para_flag.append({'para_num':para_num,'position':i})

                # 顺便添加heading_obj
                heading_exam = para_obj(type=type_name, position=pos_num, origin=text, str_=str_split, para_num=para_num,
                             from_global=len(global_obj),flag=1)
                is_heading=1
                heading_list.append(heading_exam)
            #添加正文  容易和heading的解包重复
            if is_heading==0:
                origin_=text
                global_examp=para_obj(type=type_name, position=pos_num, origin=origin_,str_=str_split,para_num=para_num,flag=1)
                global_obj.append(global_examp)
                pos_num+=1

            else: #已经被heading_obj收录
                global_examp=heading_exam
                global_obj.append(global_examp)
                pos_num+=1
    return heading_list,global_obj

trasbin=set(['','\n',' ','  ',])
h1,g1=exctract_heading(doc1.paragraphs)
h2,g2=exctract_heading(doc2.paragraphs)

docu1=[]
for i,j in enumerate(g1):
    docu1.append(j.origin)

docu2=[]
for i,j in enumerate(g2):
    docu2.append(j.origin)

# print(docu2[:10])

def extract_sen(x):
    sent=[]
    for i in range(len(x)):# 先遍历段
        j=0
        length=len(x[i])
        while j<len(x[i]):
            k=j
            while k<length and x[i][k]!='。':
                # print('k为',k,'x[0]长度为:',len(x[i]))
                k+=1
            sent.append(x[i][j:k+1])
            j=k+1
    return sent

docu1=extract_sen(docu1)
docu2=extract_sen(docu2)


from test import simhash
import jieba













# hash_list1=[]
# for i,j in enumerate(docu1):
#     j=list(jieba.cut(j))
#     hash = simhash(j)
#     hash_list1.append(hash) # hash是一个对象
# print('hash_list1:',hash_list1[:10])
#
# hash_list2=[]
# for i,j in enumerate(docu2):
#     j=list(jieba.cut(j))
#     hash = simhash(j)
#     hash_list2.append(hash) # hash是一个对象
# print('hash_list2:',hash_list2[:10])
#
# def func(a,b):
#     dis=a.hamming_distance(b)
#     return dis
#
# final_mat=[[0 for i in range(len(hash_list2))] for i in range(len(hash_list1))]  # hash1是行数
# print('hash_list1的长度:',len(hash_list1))
# print('final的长度',len(final_mat)) # 935
# for i in range(len(hash_list1)):
#     for j in range(len(hash_list2)):
#         dis=func(hash_list1[i],hash_list2[j])
#         final_mat[i][j]=dis
#         # print(dis)
# # print(final_mat)
#
# def find_min(x):
#     min_=1000
#     index=-1
#     for i,j in enumerate(x):
#         if j<min_:
#             index=i
#             min_=j
#     '''
#     min_: 最小值
#     index: 下标
#     '''
#     return min_,index
# # x=[10,6,39,9,38,1,2,7,4]
# # min_=find_min(x)
# # print('最小值:',min_)
#
# close_mat=[]
# for i,j in enumerate(hash_list1):
#     min_,index=find_min(final_mat[i])
#     close_mat.append(tuple([i,min_,index,docu1[i],docu2[index]]))
#
# print('最近链接:',close_mat)
#
# sorted_res=sorted(close_mat, key=lambda x:x[1], reverse=False)
# print('最后结果:',sorted_res)
#
#
# for i,j in enumerate(sorted_res):
#     doc1_index,dis,doc2_index,doc1,doc2=j
#     rate=hash_list1[doc1_index].dup_rate(hash_list2[doc2_index])
#     sorted_res[i]=tuple([rate,doc1_index,dis,doc2_index,doc1,doc2])
# print('rate_res是什么?',sorted_res)
#
# select_final=[]
# for i,j in enumerate(sorted_res):
#     rate,doc1_index, dis, doc2_index, doc1, doc2 = j
#     if len(doc1)>13 and len(doc2)>13:
#         select_final.append(j)
# print('达到长度要求的:',select_final)


def create_hash_obj_list(sen_list):
    hash_list = []
    for i, j in enumerate(sen_list):
        j = list(jieba.cut(j)) #生成tokens
        hash = simhash(j)
        hash_list.append(hash)
    return hash_list

def func(a,b):
    dis=a.hamming_distance(b)
    return dis

def find_min(x):
    min_=1000
    index=-1
    for i,j in enumerate(x):
        if j<min_:
            index=i
            min_=j
    '''
    min_: 最小值
    index: 下标
    '''
    return min_,index



def comp_dis_mat(hash_list1,hash_list2):
    dis_mat = [[0 for i in range(len(hash_list2))] for i in range(len(hash_list1))]  # hash1是行数
    print('hash_list1的长度:', len(hash_list1))
    print('dis_mat的长度', len(dis_mat))  # 935
    for i in range(len(hash_list1)):
        for j in range(len(hash_list2)):
            dis = func(hash_list1[i], hash_list2[j])
            dis_mat[i][j] = dis
    return dis_mat


def get_closest(hash_list1,dis_mat,docu1,docu2):
    close_list = []
    for i, j in enumerate(hash_list1):
        min_, index = find_min(dis_mat[i])
        close_list.append(tuple([i, min_, index, docu1[i], docu2[index]]))
    return close_list

# "中国水利电力物资集团有限公司(以下简称“物资集团”)是中国大唐集团公司全资子公司，注册资本10.12亿元。"
docu3=["招标人：中国水利电力物资集团有限公司","中国水利电力物资集团有限公司(以下简称“物资集团”)是中国大唐集团公司全资子公司，注册资本10.12亿元。","招标代理机构：北京国电工程招标有限公司",r"中国水利电力物资集团有限公司(以下简称“物资集团”)是中国大唐集团公司全资子公司，注册资本10.12亿元。主要经营招标代理、进出口代理、管道及油料供应、工程技术咨询、备品配件、设备监理、工程物资管理、碳资产开发、安全性评价、煤化工综合服务、电子商务等业务。"]

x=['2345',"中国水利电力物资集团有限公司(以下简称“物资集团”)是中国大唐集团公司全资子公司，注册资本10.12亿元。",'123。']
x=x[1]
x = list(jieba.cut(x)) #生成tokens
hash = simhash(x)
print('第一次哈希:',hash.hash)

x=r"中国水利电力物资集团有限公司(以下简称“物资集团”)是中国大唐集团公司全资子公司，注册资本10.12亿元。"
x = list(jieba.cut(x)) #生成tokens
hash = simhash(x)
print('第二次哈希:',hash.hash)


def find_match(docu1,docu2,docu3):
    '''
    docu1:一个list，元素为一句话
    '''
    hash_list1=create_hash_obj_list(docu1)
    hash_list2=create_hash_obj_list(docu2)
    hash_list3 = create_hash_obj_list(docu3)

    dis_mat=comp_dis_mat(hash_list1,hash_list2)
    dis_mat2 = comp_dis_mat(hash_list1, hash_list3) # 查找与模板的相似
    # print('dis_mat2是什么?',dis_mat2)


    print('docu1的第22个:',docu1[22])
    print('是否相等:',docu1[22]=="中国水利电力物资集团有限公司(以下简称“物资集团”)是中国大唐集团公司全资子公司，注册资本10.12亿元。")
    print('dis_mat的第22个:',dis_mat[22])
    print('dis_mat2的第22个',dis_mat2[22])

    close_list=get_closest(hash_list1, dis_mat, docu1, docu2)
    close_list2 = get_closest(hash_list1, dis_mat2, docu1, docu3)

    print('close_list长度:',len(close_list))
    print('close_list2长度:', len(close_list2))

    tichu_list=[0 for i in range(len(close_list2))]
    for i,j in enumerate(close_list2):
        doc1_index, dis, doc2_index, doc1, doc2 = j #解包
        if dis<=4:
            tichu_list[i]=1 #在模板中有，点亮，表明要除去
    for i,j in enumerate(tichu_list):
        if j ==1:
            print('剔除结果:',docu1[i])

    #去重最好也在这里做，不然等下又要重新算hash值

    # sorted_list = sorted(close_list, key=lambda x: x[1], reverse=False)
    # sorted_list2 = sorted(close_list2, key=lambda x: x[1], reverse=False)

    # print('sorted_list的前20个:',sorted_list[:20])
    # print('sorted_list2的前20个:', sorted_list2[:20])

    #计算重复率
    no_docu3_list=[]
    for i, j in enumerate(close_list):
        if tichu_list[i]==0: #不剔除的才计算重复率
            doc1_index, dis, doc2_index, doc1, doc2 = j
            rate = hash_list1[doc1_index].dup_rate(hash_list2[doc2_index])
            # sorted_list[i] = tuple([rate, doc1_index, dis, doc2_index, doc1, doc2])
            no_docu3_list.append(tuple([rate, doc1_index, dis, doc2_index, doc1, doc2]))
    print('no_docu3_list是什么?', no_docu3_list)
    print('no_docu3_list的长度是什么',len(no_docu3_list))
    sorted_list = sorted(no_docu3_list, key=lambda x: x[2], reverse=False)

    select_final = []
    sen_count=0
    for i, j in enumerate(sorted_list):
        if sen_count>20: #取出最接近的20个
            break
        rate, doc1_index, dis, doc2_index, doc1, doc2 = j
        if len(doc1) > 13 and len(doc2) > 13:
            select_final.append(j)
            sen_count+=1
    print('达到长度要求的:', select_final)


find_match(docu1,docu2,docu3)



