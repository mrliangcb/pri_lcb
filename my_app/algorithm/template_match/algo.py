
import docx
import collections
from collections import namedtuple as nt

para_obj =nt('paragraph', ['type', 'position', 'origin','str_','flag','test']) # flag和test怎么用
para_obj.__new__.__defaults__ = ('para',None, None,None,None,None)

def extract_4para(doc_file):  #主要用于处理模板文章
    save_flag = 0
    heading4_obj = []
    global_obj=[]

    for i,paragraph in enumerate(doc_file.paragraphs):
        tem_obj=None
        if paragraph.text!='':
            if paragraph.text.strip('\n').startswith('第四章'):
                save_flag = 1
            if paragraph.text.strip('\n').startswith('第五章'):
                save_flag = 0
            # print('模板的:',paragraph.style.name,paragraph.text)
            # tem_obj = para_obj(type=paragraph.style.name, position=i, origin=paragraph.text,str_=paragraph.text.strip().split(' ')[-1])  # type=normal/Heading *
            # print('是否保存:',paragraph.style.name,paragraph.text,save_flag)
            # save_flag=1
            if save_flag == 1:
                if paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('List Paragraph'):#这个para是heading
                        x=para_obj(type=paragraph.style.name, position=i, origin=paragraph.text,str_=paragraph.text.strip().split(' ')[-1]) #1.64s
                    # x = para_obj(type=paragraph.style.name, position=i) #ss1.5s
                    # x={'type':paragraph.style.name,'position':i,'obj':paragraph} # 1.67s
                        heading4_obj.append(x)
            # global_obj.append(tem_obj)#要存这个的话好久
    return global_obj,heading4_obj


def exctract_heading(para_list):
    heading_list=[]
    for i,para in enumerate(para_list):
        if para.text != '':
        # print(para.style.name)
        # print('第{}段:'.format(i),para.style.name,para.text)
            if para.style.name.startswith('Heading') or para.style.name.startswith('List Paragraph'):
                x=para_obj(type=para.style.name, position=i, origin=para.text,str_=para.text.strip().split(' ')[-1])
                heading_list.append(x)
    return heading_list

class processer():
    def read_doc(self,path):#path or io
        x = docx.Document(path)
        return x

class Solution:
    def lengthOfLIS(self, nums):
        print('nums是什么',nums)
        dp = [1 for i in range(len(nums))]  # 用于存储每一个元素处的最大序列的长度
        dp2=[[i] for i in range(len(nums))] #存储序列   #大家默认只取自己
        print('dp2是什么?',dp2)
        n = len(nums)
        max_len = 1
        max_id=0
        for i in range(n):
            if nums[i].test >= 0:
                tmp = 1 #因为最小可以取i自己
                tmp2=[i] #首先取自己
                for j in range(0,i):
                    if nums[j].test >= 0:
                        if nums[j].test<nums[i].test:#表明i可以用自己的，也可以自己接上j的 now
                            # tmp = max(tmp,1+dp[j])
                            now=1+dp[j]
                            if tmp>now:#tmp是最新i的
                                pass
                                # tmp 不变
                            else:
                                tmp=now
                                tmp2=dp2[j]+[i]#j的地方连接上i 下标
                dp[i]=tmp
                dp2[i]=tmp2 #这个是装下标的
                if max_len<tmp:
                    max_len=tmp
                    max_id =i
        print('最后dp2:',dp2)
        return max_len,dp2[max_id] #dp2是下标   如果喜test都是-2，那会返回0

def make_seq(x,y):#x list  y dic   投标文档 参照字典，重做下标    有可能重复标题
    seq_=[-2 for i in range(len(x))]
    print('初始化的seq_:',len(seq_),seq_)
    for i,j in enumerate(x):
        temp=y.get(j.str_)  # 模板字典的位置
        if temp!=None:
            x[i] = x[i]._replace(test=temp) # 与flag区别    test是暂时用于放template的对应下标
            seq_[i]=x[i]
        else:
            x[i] = x[i]._replace(test=-2)
            seq_[i] =x[i] #没找到
    return seq_
def find_best_match(heading4_target_obj_list,source_heading_obj_list):
    all_heading1_dic = {}
    all_heading1_list = []
    for i, j in enumerate(heading4_target_obj_list):
        if not all_heading1_dic.get(j.str_):  # 不重复的标题
            all_heading1_dic[j.str_] = i
            all_heading1_list.append(j.str_)   #如果重复了，算是有但错位

    print('source_heading_obj_list?',source_heading_obj_list)

    seq = make_seq(source_heading_obj_list, all_heading1_dic) #shape=source   元素为template的下标
    all_heading1_set = set(all_heading1_list) #template集合

    source_heading_list_str = []
    for i, j in enumerate(source_heading_obj_list):
        source_heading_list_str.append(j.str_)
    source_heading_set_str = set(source_heading_list_str) #source集合


    print('seq是什么:?',seq) # [paragraph(type='Heading 5', position=633, origin='', str_='', flag=None, test=-2)]
    exam = Solution()
    result = exam.lengthOfLIS(seq) #seq要求包含 tem下表，也要有type
    print('最长公共结果result',result) # (1, 0)

    #左边
    # 解决序列对上的，级别是否对上
    flag_left = [-2 for i in range(len(all_heading1_list))]
    flag_right = [-2 for i, _ in enumerate(source_heading_list_str)]


    for i, j in enumerate(result[1]):  # seq的最长公共子序列下标
        template_index = seq[j].test #对应到template下标  seq是source的对象list
        if seq[j].type == heading4_target_obj_list[template_index].type:
            tem_flag = 1
        else:
            tem_flag = -4
        flag_left[template_index] = tem_flag
        flag_right[j] = tem_flag

    # 左边 解决是否存在的问题
    for i, j in enumerate(flag_left):
        if j == -2:  # 只检查-2的情况
            if all_heading1_list[i] in source_heading_set_str:  # 在但位置不对
                flag_left[i] = -3

    # 右边 解决是否存在的问题
    for i,j in enumerate(flag_right):
        if j==-2:#只检查-2的情况
            if source_heading_list_str[i] in all_heading1_set: #在但位置不对
                flag_right[i]=-3

    # 做返回的obj
    for i, j in enumerate(flag_left):
        heading4_target_obj_list[i] = heading4_target_obj_list[i]._replace(flag=j)
    for i, j in enumerate(flag_right):
        source_heading_obj_list[i] = source_heading_obj_list[i]._replace(flag=j)

    # 测试用显示
    left_print = []
    for i, j in enumerate(flag_left):
        left_print.append([j, heading4_target_obj_list[i].str_])
    right_print = []
    for i, j in enumerate(flag_right):
        right_print.append([j, source_heading_obj_list[i].str_])
    print('左边',left_print)
    print('右边:', right_print)

    return heading4_target_obj_list,source_heading_obj_list
import time
def main(source,template):
    process_time=time.time()
    procer = processer()
    template_doc = procer.read_doc(template)
    print('解析时间1.1:', time.time() - process_time)

    process_time = time.time()
    global_obj_target_obj_list, heading4_target_obj_list = extract_4para(template_doc)
    print('解析时间1.2:', time.time() - process_time)


    process_time = time.time()
    source_file = procer.read_doc(source)
    source_heading_obj_list = exctract_heading(source_file.paragraphs)
    print('解析时间2:', time.time() - process_time)

    mat_time=time.time()
    a,b=find_best_match(heading4_target_obj_list,source_heading_obj_list)
    print('匹配time:',time.time()-mat_time)
    return a,b

if __name__ == '__main__':
    base=r'D:\lcb_note\code\Program\10月项目\my_docx'
    path1 = base+r'\招标文件 CWEME-1911ZSWZ-2J039 基于NLP的商务文本数据清洗关键技术研究项目-2019年12月中国水利电力物资集团有限公司项目（第三版终版）.docx'
    path2 = base+r'\基于NLP的商务文本数据清洗关键技术研究项目合同+-+-打印版.docx'
    procer = processer()
    doc_1 = procer.read_doc(path1)

    global_obj_target_obj_list, heading4_target_obj_list = extract_4para(doc_1)
    doc2_file = procer.read_doc(path2)
    source_heading_obj_list = exctract_heading(doc2_file.paragraphs)
    a,b=find_best_match(heading4_target_obj_list,source_heading_obj_list)
    print(a)
    print(b)
