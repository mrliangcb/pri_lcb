import docx
import collections
from collections import namedtuple as nt

para_obj = nt('paragraph',
              ['type', 'position', 'origin', 'str_', 'flag', 'test', 'para_num', 'from_global','title'])  # flag和test怎么用  flag表示是否正确
para_obj.__new__.__defaults__ = ('para', None, None, None, None, None, None,None,None)

# def extract_4para(doc_file):  #主要用于处理模板文章
#     save_flag = 0
#     heading4_obj = []
#     global_obj=[]
#
#     for i,paragraph in enumerate(doc_file.paragraphs):
#         tem_obj=None
#         if paragraph.text!='':
#             if paragraph.text.strip('\n').startswith('第四章'):
#                 save_flag = 1
#             if paragraph.text.strip('\n').startswith('第五章'):
#                 save_flag = 0
#             if save_flag == 1:
#                 if paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('List Paragraph'):#这个para是heading
#                         x=para_obj(type=paragraph.style.name, position=i, origin=paragraph.text,str_=paragraph.text.strip().split(' ')[-1]) #1.64s
#                     # x = para_obj(type=paragraph.style.name, position=i) #ss1.5s
#                     # x={'type':paragraph.style.name,'position':i,'obj':paragraph} # 1.67s
#                         heading4_obj.append(x)
#             # global_obj.append(tem_obj)#要存这个的话好久
#     return global_obj,heading4_obj


trasbin = set(['', '\n', ' ', '  ', ])

import re


def exctract_heading(para_list):
    heading_list = []
    para_num = -1
    global_obj = []
    pos_num = 0
    for i, para in enumerate(para_list):
        text = para.text.strip()
        if (text not in trasbin) and text:
            # 有两个解析，等判空之后再做，减少一点时间，剪枝
            type_name = para.style.name
            str_split = text.replace(' ','') #去掉空格  1.2标题内容  1.2 标题内容
            is_heading = 0
            # 校正章号
            title=0
            if type_name.startswith('Heading'):
                # if (not para.style.name.startswith('Normal')) and (not para.style.name.startswith('normal')):
                ptr = r'第(.*?)章'  # 非贪心
                ptr = r'(第[0-9一二三四五六七八九十]+章[\s]*.*)'
                res = re.match(ptr, text)
                if res:#大标题
                    title=1
                # result = re.findall(ptr, text)
                # if result and result[0] != '':
                #     para_num += 1
                    # 这是一个章标题
                    # para_flag.append({'para_num':para_num,'position':i})

                # 顺便添加heading_obj
                heading_exam = para_obj(type=type_name, position=pos_num, origin=text, str_=str_split,from_global=len(global_obj), flag=1,title=title)# para_num=para_num,
                is_heading = 1
                heading_list.append(heading_exam)
            # 添加正文  容易和heading的解包重复
            if is_heading == 0:
                origin_ = text
                global_examp = para_obj(type=type_name, position=pos_num, origin=origin_, str_=str_split,flag=2)# para_num=para_num
                global_obj.append(global_examp)
                pos_num += 1

            else:  # 已经被heading_obj收录
                global_examp = heading_exam
                global_obj.append(global_examp)
                pos_num += 1
    return heading_list, global_obj


class processer():
    def read_doc(self, path):  # path or io
        x = docx.Document(path)
        return x


class Solution:
    def lengthOfLIS(self, nums):
        dp = [1 for i in range(len(nums))]  # 用于存储每一个元素处的最大序列的长度
        dp2 = [[i] for i in range(len(nums))]  # 存储序列   #大家默认只取自己

        n = len(nums)
        max_len = 1
        max_id = 0
        for i in range(n):
            if nums[i].test >= 0:
                tmp = 1  # 因为最小可以取i自己
                tmp2 = [i]  # 首先取自己
                for j in range(0, i):
                    if nums[j].test >= 0:
                        if nums[j].test < nums[i].test:  # 表明i可以用自己的，也可以自己接上j的 now
                            # tmp = max(tmp,1+dp[j])
                            now = 1 + dp[j]
                            if tmp > now:  # tmp是最新i的
                                pass
                                # tmp 不变
                            else:
                                tmp = now
                                tmp2 = dp2[j] + [i]  # j的地方连接上i 下标
                dp[i] = tmp
                dp2[i] = tmp2  # 这个是装下标的
                if max_len < tmp:
                    max_len = tmp
                    max_id = i
        return max_len, dp2[max_id]  # dp2是下标   如果喜test都是-2，那会返回0


def make_seq(x, y):  # x list  y dic   投标文档 参照字典，重做下标    有可能重复标题
    seq_ = [-2 for i in range(len(x))]

    for i, j in enumerate(x):
        temp = y.get(j.str_)  # 模板字典的位置
        if temp != None:
            x[i] = x[i]._replace(test=temp)  # 与flag区别    test是暂时用于放template的对应下标   #这个操作会改变原来的x的test   temp就是在标题对象中的序号
            seq_[i] = x[i]
        else:
            x[i] = x[i]._replace(test=-2)  # 找不到下标
            seq_[i] = x[i]  # 没找到
    return seq_


def find_best_match(heading4_target_obj_list, tem_global_obj_list, source_heading_obj_list, source_global_obj_list):
    print('模板标题对象:', heading4_target_obj_list, len(heading4_target_obj_list))
    print('source标题对象:', source_heading_obj_list, len(source_heading_obj_list))
    # print('source全文标题对象:', source_global_obj_list, len(source_global_obj_list))

    all_heading1_dic = {}
    all_heading1_list = []
    over_heading = 0
    for i, j in enumerate(heading4_target_obj_list):
        if not all_heading1_dic.get(j.str_):  # 不重复的标题   将模板标题的str放到  字典
            all_heading1_dic[j.str_] = i  # str -> i   i是标题对象的序号
            # 如果重复了，算是有但错位
        else:
            over_heading = 1  # 重复标题出现
        all_heading1_list.append(j.str_)  # 含重复标题

    print('是否有重复标题:', over_heading == 1)

    seq = make_seq(source_heading_obj_list, all_heading1_dic)  # shape=source   元素为template的下标
    # seq是source标题对象，test被赋予了 模板标题的序号
    # 按道理来说source_heading_obj和seq是一样的
    all_heading1_set = set(all_heading1_list)  # 有重复标题的时候，set变小
    print('all_heading1_list与heading4_target_obj_list的长度:', len(all_heading1_list), len(heading4_target_obj_list))

    source_heading_list_str = []
    for i, j in enumerate(source_heading_obj_list):
        source_heading_list_str.append(j.str_)  # source的内容做个集合
    source_heading_set_str = set(source_heading_list_str)  # source集合

    exam = Solution()
    result = exam.lengthOfLIS(seq)  # seq要求包含 tem下表，也要有type
    print('最长公共结果result', result)  # (1, 0)

    # 左边
    # 解决序列对上的，级别是否对上
    flag_left = [-2 for i in
                 range(len(heading4_target_obj_list))]  # 如果出现重复标题 heading4_target_obj_list 可能大于这个 all_heading1_list
    flag_right = [-2 for i, _ in enumerate(source_heading_list_str)]

    for i, j in enumerate(result[1]):  # seq的最长公共子序列下标
        template_index = seq[j].test  # j是seq里面的下表  .test是模板题目的下标        对应到template下标  seq是source的对象list
        if seq[j].type == heading4_target_obj_list[template_index].type:
            tem_flag = 1
        else:
            tem_flag = 1  # -4
        try:
            flag_left[template_index] = tem_flag
        except:
            print('flag_left[template_index]越界了：', len(flag_left), flag_left, template_index)
        flag_right[j] = tem_flag

    # 左边 解决是否存在的问题
    for i, j in enumerate(flag_left):
        if j == -2:  # 只检查-2的情况
            if all_heading1_list[i] in source_heading_set_str:
                flag_left[i] = -3

    print('flag_left是什么:', flag_left)
    print('flag_right是什么:', flag_right)
    # 右边 解决是否存在的问题
    for i, j in enumerate(flag_right):
        if j == -2:  # 只检查-2的情况
            if source_heading_list_str[i] in all_heading1_set:  # 在但位置不对
                flag_right[i] = -3

    # 做返回的obj
    for i, j in enumerate(flag_left):
        heading4_target_obj_list[i] = heading4_target_obj_list[i]._replace(flag=j)
        global_index = heading4_target_obj_list[i].from_global
        tem_global_obj_list[global_index] = tem_global_obj_list[global_index]._replace(flag=j)

    # print('flag_right:',len(flag_right))
    # print('source_heading_obj_list:', len(source_heading_obj_list),source_heading_obj_list)
    # print('source_global_obj_list:',source_global_obj_list)

    # flag_right 和 source_heading_obj_list 都是 长度30
    for i, j in enumerate(flag_right):
        source_heading_obj_list[i] = source_heading_obj_list[i]._replace(flag=j)

        global_index = source_heading_obj_list[i].from_global
        source_global_obj_list[global_index] = source_global_obj_list[global_index]._replace(flag=j)

    # 测试用显示
    left_print = []
    for i, j in enumerate(flag_left):
        left_print.append([j, heading4_target_obj_list[i].str_])
    right_print = []
    for i, j in enumerate(flag_right):
        right_print.append([j, source_heading_obj_list[i].str_])

    return heading4_target_obj_list, source_heading_obj_list, source_global_obj_list


import time

from collections import Counter


def get_muban(doc1_global_para, source_heading_obj_list):
    doc1_list = []
    doc1_dic = {}
    doc1_set = 0
    para_obj_dict = {}
    for i, j in enumerate(doc1_global_para):  # 换一种数据结构，容易搜索
        if not doc1_dic.get(j.str_, None):
            doc1_dic[j.str_] = j.para_num
        if para_obj_dict.get(j.para_num, None):  # 如果存在
            para_obj_dict[j.para_num].append(j)
        else:  # 不存在，就新建
            para_obj_dict[j.para_num] = [j]

    doc2_para_num = []
    for i, j in enumerate(source_heading_obj_list):
        if doc1_dic.get(j.str_, None) != None:
            doc2_para_num.append(doc1_dic[j.str_])

    # para_num_list=list(dict(y).keys())
    # tem_para=Counter(doc2_para_num).most_common(1)[0][0]

    tem_para = Counter(doc2_para_num).most_common(1)
    print('最匹配的段counter:', Counter(doc2_para_num))
    if tem_para != []:
        tem_para = tem_para[0][0]
        result = para_obj_dict[tem_para]  # 取出最匹配的那个段
    else:
        result = 0
    print('最匹配的模板段:', result)
    return result


trasbin = set(['', '\n', ' ', '  ', ])


def extract_doc_heading(para_list: list):
    heading_list = []
    para_num = -1
    global_obj = []
    pos_num = 0
    for i, para in enumerate(para_list):
        # print('第{}个para是什么:{}'.format(i,para))
        text = para['text'].strip()
        type_name = para['style']
        str_split = text.replace(' ','') # text.split(' ')[-1]
        title_level = 0
        if (text not in trasbin) and text:  # 去除空段
            is_heading = 0
            # 校正标章号  para_num应该是chapter_num

            if type_name.startswith('标题'):
                # ptr = r'第(.*?)章'  # 非贪心
                # result = re.findall(ptr, text)
                # if result and result[0] != '':  # 是一个章的标志位
                #     para_num += 1
                ## 顺便添加heading_obj
                ptr = r'(第[0-9一二三四五六七八九十]+章[\s]*.*)'
                res = re.match(ptr, text)
                if res:  # 大标题
                    title_level = 1

                heading_exam = para_obj(type=type_name, position=pos_num, origin=text, str_=str_split,from_global=len(global_obj), flag=1,title=title_level)# para_num=para_num,
                is_heading = 1
                heading_list.append(heading_exam)
            # 保存正文
            if is_heading == 0:
                origin_ = text
                global_examp = para_obj(type=type_name, position=pos_num, origin=origin_, str_=str_split,flag=2) # para_num=para_num,
                global_obj.append(global_examp)
                pos_num += 1
            else:  # 已经被heading_obj收录
                global_examp = heading_exam
                global_obj.append(global_examp)
                pos_num += 1
    return heading_list, global_obj


def main(source_file, template_doc, source_isdoc, tem_isdoc):
    process_time = time.time()
    # procer = processer()
    # template_doc = procer.read_doc(template)
    if tem_isdoc == 0:
        print('tem进行docx解码')
        tem_heading_obj_list, tem_global_obj_list = exctract_heading(template_doc.paragraphs)
        # print('模板题目的解析:',tem_heading_obj_list)
        # print('模板全文的解析:',tem_global_obj_list)

    else:  # 是doc文件
        print('tem进行doc解码')
        print('template_doc的例子:', template_doc[:2])
        tem_heading_obj_list, tem_global_obj_list = extract_doc_heading(template_doc)  #

    print('解析时间1.1:', time.time() - process_time)

    process_time = time.time()
    # source_file = procer.read_doc(source)
    if source_isdoc == 0:
        print('sour进行docx解码')
        source_heading_obj_list, source_global_obj_list = exctract_heading(source_file.paragraphs)
        # print('source题目的解析:', source_heading_obj_list)
    else:
        print('sour进行doc解码')
        source_heading_obj_list, source_global_obj_list = extract_doc_heading(source_file)
    # source_global_list_obj = extract_global(source_file.paragraphs)
    print('解析时间2:', time.time() - process_time)

    tem_heading_str_list=[i.str_ for i in tem_heading_obj_list] #会不会出现重复的问题
    tem_heading_str_set=set(tem_heading_str_list)

    source_heading_str_list = [i.str_ for i in source_heading_obj_list]  # 会不会出现重复的问题
    source_heading_str_set = set(source_heading_str_list)

    flag_tem=[0 for i in range(len(tem_heading_str_list))]


    print('tem中没有重复标题:',len(tem_heading_str_set)==len(tem_heading_str_list)) #重复的概率比较小， 位置编号+内容

    for i,j in enumerate(tem_heading_str_list):
        # print('i是什么?',i)
        # print('source_heading_obj_list的长度:',len(source_heading_obj_list))
        # print('tem_heading_obj_list的长度:',len(tem_heading_obj_list))
        if j in source_heading_str_set:
            # tem的这个标题正确 (有，并且第几章  2.1这个位置也正确)
            tem_heading_obj_list[i] = tem_heading_obj_list[i]._replace(flag=1)
            global_index = tem_heading_obj_list[i].from_global
            tem_global_obj_list[global_index]=tem_global_obj_list[global_index]._replace(flag=1)
            flag_tem[i]=1
        else:
            tem_heading_obj_list[i]=tem_heading_obj_list[i]._replace(flag=-2) #
            global_index = tem_heading_obj_list[i].from_global
            tem_global_obj_list[global_index] = tem_global_obj_list[global_index]._replace(flag=-2)
            flag_tem[i] = 0

    flag_source = [0 for i in range(len(source_heading_str_list))]
    print('source中没有重复标题:', len(source_heading_str_set) == len(source_heading_str_list))  # 重复的概率比较小， 位置编号+内容

    for i,j in enumerate(source_heading_str_list):
        if j in tem_heading_str_set:
            # tem的这个标题正确 (有，并且第几章  2.1这个位置也正确)
            source_heading_obj_list[i] = source_heading_obj_list[i]._replace(flag=1)
            global_index = source_heading_obj_list[i].from_global
            source_global_obj_list[global_index]=source_global_obj_list[global_index]._replace(flag=1)
            flag_source[i]=1
        else:
            source_heading_obj_list[i] = source_heading_obj_list[i]._replace(flag=-2)
            global_index = source_heading_obj_list[i].from_global
            source_global_obj_list[global_index] = source_global_obj_list[global_index]._replace(flag=-2)
            flag_source[i] = 0

    template_select_obj_list = tem_heading_obj_list

    print('source_heading_obj_list:',source_heading_obj_list)
    print('tem_heading_obj_list:',tem_heading_obj_list)


    # time_find_tem=time.time()
    # template_select_obj_list = get_muban(tem_heading_obj_list, source_heading_obj_list)
    # print('找模板时间:',time.time()-time_find_tem)

    # mat_time = time.time()
    # tem_heading_match, source_heading, source_global_obj = find_best_match(template_select_obj_list,
    #                                                                        tem_global_obj_list, source_heading_obj_list,
    #                                                                        source_global_obj_list)
    # print('计算最长匹配子串时间:', time.time() - mat_time)

    left_2 = 0
    b_t_1=0 #缺失的大标题
    s_t_1=0 #缺失的小标题
    for i, j in enumerate(flag_tem):
        if j == 0:
            if tem_heading_obj_list[i].title ==1: #大标题
                b_t_1+=1
            else:
                s_t_1 +=1

    b_t_2 = 0 #正确的大标题1
    s_t_2 = 0 #正确的小标题
    r_b=0 # 右边所有的大标题
    r_s=0 # 右边所有的小标题
    correct_heading = 0
    for i, j in enumerate(flag_source):
        if j == 1:
            if source_heading_obj_list[i].title == 1:  # 大标题
                b_t_2+=1
            else:
                s_t_2+=1

    for i, j in enumerate(source_heading_obj_list):
        if j.title == 1:
            r_b+=1
        else:
            r_s+=1

    print('正确大标题:',b_t_2)
    print('正确小标题:', s_t_2)

    print('r_b:{},r_s:{}'.format(r_b,r_s))
    print('b_t_1:{},b_t_2:{}'.format(b_t_1, b_t_2))

    # print('计算匹配值:{}/{}'.format(correct_heading, len(flag_source) + left_2)) # (source中的正确标题)/(source全部标题+tem的缺失标题)
    try:
        b_score = 0.65 * (b_t_2 / (r_b + b_t_1))
    except:
        b_score = 0
    try:
        s_score = 0.35 * (s_t_2 / (r_s + s_t_1))
    except:
        s_score = 0
    global_score = b_score + s_score

    print('global_score是多少?', global_score)

    '''
    tem_heading_match:left
    source_heading:right
    tem_global_obj_list:tem_global_list_obj  应该是按照原样返回的，  flag全部是1
    source_global_obj: 修改了flag的，可以变红色  相当于将right放回到global
    match_rate_head： 标题匹配度

    '''
    return flag_tem, flag_source, tem_global_obj_list, source_global_obj_list, global_score


def main2():
    pass


if __name__ == '__main__':
    base = r'D:\lcb_note\code\Program\10月项目\my_docx'
    path1 = base + r'\招标文件 CWEME-1911ZSWZ-2J039 基于NLP的商务文本数据清洗关键技术研究项目-2019年12月中国水利电力物资集团有限公司项目（第三版终版）.docx'
    path2 = base + r'\基于NLP的商务文本数据清洗关键技术研究项目合同+-+-打印版.docx'
    procer = processer()
    doc_1 = procer.read_doc(path1)
    doc2_file = procer.read_doc(path2)
    # global_obj_target_obj_list, heading4_target_obj_list = extract_4para(doc_1)
    doc1_global_para, tem_global_obj = exctract_heading(doc_1.paragraphs)

    source_heading_obj_list, source_global_obj = exctract_heading(doc2_file.paragraphs)
    source_global_list_obj = extract_global(doc2_file.paragraphs)

    template_obj_list = get_muban(doc1_global_para, source_heading_obj_list)
    a, b, source_global_obj_list = find_best_match(template_obj_list, source_heading_obj_list, source_global_obj)















