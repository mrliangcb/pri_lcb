from flask import jsonify,Blueprint,request,render_template
from my_app.forms import check_args_validation,vali_check_match1
from . import web
from my_app.algorithm.dup_check_algo import check_str
from my_app.algorithm.duan_winnowing3 import paragraph_winnowing
import logging
import requests
import json
# https://blog.csdn.net/weixin_30773135/article/details/97342082
# file = open("demo.log", encoding="utf-8", mode="a")
logging.basicConfig(format='%(asctime)s %(filename)s %(levelname)s %(message)s',datefmt='%a %d %b %Y %H:%M:%S',filename='demo.log',level=logging.DEBUG)# stream=file,
logging.warning("warning")
from io import BytesIO
import docx
import json
from my_app.algorithm.simhash.my_simhash2 import start_jieba,sim_main
y = start_jieba()
# -*- coding: utf-8 -*-
import codecs,sys
# sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())
# sys.stdout.write("Your content....")
 #把注册好的蓝图拿来用
# if sys.stdout.encoding != 'UTF-8':
#     sys.stdout = codecs.getwriter('utf-8')(sys.stdout, 'strict')


@web.route('/NLP/Algorithm/base/dup_check/winnowing3', methods=['POST','GET'])
def dup_check3():
    # args_dic = request.args  # 这个是不可变字典，如果转成普通字典
    print('now winnowing2')
    args_dic=request.form.to_dict()

    try:#网页传参模式
        dic=request.args.to_dict()
        doc1=dic['doc1']
        doc2=dic['doc2']
        a, b = check_args_validation(dic)
    except:# postman   form-data或者www-form模式
        dic=request.form.to_dict()
        doc1=dic['doc1']
        doc2=dic['doc2']
        a, b = check_args_validation(dic)
    if not a:
        return b


    time_, result = check_str(str(doc1), str(doc2), k=13)
    similiarize, dup_text, dup_dic, doc1_str, doc2_str, doc1_wrap, doc2_wrap = result

    result1 = str(similiarize)
    result2 = dup_dic
    result3 = render_template('testHtml.html', name1='doc1', name2='doc2', time=time_, dup_check=similiarize,
                              dup_text=dup_text, dup_dic=dup_dic, doc1_str=doc1_str, doc2_str=doc2_str,
                              doc1_wrap=doc1_wrap, doc2_group_=doc2_wrap)
    result4 = render_template('add_href_doc1.html', doc1_wrap=doc1_wrap, doc1_str=doc1_str)
    result5 = render_template('add_href_doc2.html', doc2_group_=doc2_wrap, doc2_str=doc2_str)

    print('dup_text：',result2)
    result6 = time_
    result7 = doc1_str
    result = doc2_str
    # return result4

    result_dic = {'dup_rate': result1,
                  'source_label': result4,
                  'doc2_label': result5}
    return jsonify(result_dic)
    # return jsonify([result1, result2, result3, result4, result5, result6])


    # return result4


# http://127.0.0.1:5002/NLP/Algorithm/base/dup_check/winnowing?doc1=中双方均希望对本协议所述保密资料及信息予以有效保护&doc2=双方均希望对本协议所述保密资料及信息予以有效保护两
# 0.96

# http://127.0.0.1:50000/NLP/Algorithm/base/dup_check/winnowing?doc1=%E6%9C%80%E8%BF%91%E7%9A%84NBA%EF%BC%8C%E6%B2%A1%E6%9C%89%E6%AF%94%E8%B5%9B%EF%BC%8C%E4%BD%86%E4%BA%8B%E6%83%85%E8%BF%98%E6%98%AF%E4%B8%8D%E5%B0%91%E7%9A%84%E3%80%82%E6%AF%94%E5%A6%82%EF%BC%8C%E6%A0%B9%E6%8D%AE%E4%B8%80%E4%BA%9B%E5%AA%92%E4%BD%93%E7%9A%84%E6%8A%A5%E9%81%93%EF%BC%8C%E6%9F%90%E7%9F%A5%E5%90%8D%E7%9A%84%E5%80%92%E9%9C%89%E6%80%BB%E7%BB%8F%E7%90%86%EF%BC%8C%E6%9C%80%E7%BB%88%E8%BF%98%E6%98%AF%E9%80%89%E6%8B%A9%E4%BA%86%E8%BE%9E%E8%81%8C%E3%80%82\n%E8%BF%99%E6%A0%B7%E7%9A%84%E6%80%BB%E7%BB%8F%E7%90%86%EF%BC%8C%E6%9D%A5%E8%87%AA%E4%BA%8E%E9%A9%AC%E8%B5%9B%E5%85%8B%E9%98%9F%EF%BC%8C%E5%A6%82%E4%BB%8A%EF%BC%8C%E6%88%96%E8%AE%B8%E5%8F%AF%E4%BB%A5%E5%8F%AB%E7%81%AB%E7%AE%AD%E9%98%9F%E4%BA%86%E3%80%82&doc2=
#

def uniform_(x):#不管输入是多少段，规整后每段至少有一个''   减少[[]]    这种情况
    for i in range(len(x)):
        if not x:#为空
            pass



def clear(x:list):#至少是''  输入是多个段
    temp = []
    transbin=set(['','\n',' ','\t'])

    for i in range(len(x)):  #遍历段
        if x[i] in transbin or (not x[i]):
            pass
        else:
            temp.append(x[i])
    if temp==[]:
        return ['']  #至少是['']
    return temp  #一维list

def source_dup_dic(result_str):#
    result_str_plus=[]
    temp=[]
    for i in range(len(result_str)):
        for j in range(0,len(result_str[i])):
            if result_str[i][j]!='':
                temp.append(result_str[i][j])
                if j==len(result_str[i])-1:
                    if temp:
                        if len(temp)>=13:
                            result_str_plus.append(''.join(temp))
                        temp = []
            else:
                if temp:
                    if len(temp) >= 13:
                        result_str_plus.append(''.join(temp))
                    temp=[]
    if not result_str_plus:result_str_plus.append('')
    source_dup_dic={}
    for k,v in enumerate(result_str_plus):
        source_dup_dic[str(k)]=v
    return source_dup_dic





def search_dot_2dec(x,num1,num2):#根据两个位置寻找前后句号
    s = 0
    e = len(x)
    # print('num1和num2',x[num1],x[num2])
    for i in range(num1,-1,-1):
        if x[i]=='。' or (num1-i)>50:
            s=i+1
            # print('找到s=句号',i,x[i])
            break

    for i in range(num2,len(x)):
        if x[i]=='。'or (i-num2)>100:
            e=i+1
            break
    return x[s:e]

def my_split(x:str)->str:
    # 先查看转义还是非转义
    split_flag='\n'
    if split_flag in x:
        pass
    else:
        split_flag=r'\n'

    x = x.split(split_flag)  # ['']  ['','','']
    x=clear(x) # ['']  ['','','']
    x_fenduan=x
    duandian = []
    for i, j in enumerate(x):
        duandian.append(len(j)) #

    # 断点递增
    for i in range(1, len(duandian)):
        duandian[i] += duandian[i - 1]

    x=''.join(x) # '123\n123'  变成'123123'


    return x_fenduan,x,duandian

import time

def get_key_data(request,key,key_ok=0):
    try:#尝试挖出参数
        try:
            dic=request.args.to_dict() #
            content=dic[key]
            key_ok = 1
        except:
            dic=request.form.to_dict() #
            content=dic[key]
            key_ok = 1
    except:
        pass
    return key_ok,content,dic

def list_model(doc1_wrap,doc2_wrap,source,target):
    doc2_wrap_dic = {}
    for i_ in range(len(doc2_wrap)):
        for j_ in range(len(doc2_wrap[i_])):
            g_, s_, e_ = doc2_wrap[i_][j_]
            if g_ != -1 and not doc2_wrap_dic.get(g_):
                doc2_wrap_dic[g_] = tuple([i_, s_, e_])

    source_target_list = []  #列表模式
    for i in range(len(doc1_wrap)):
        for j in range(len(doc1_wrap[i])):

            group_,s,e=doc1_wrap[i][j]
            if group_!=-1:
                tem_group=group_
                # print('i是什么:',i)
                # print('group_,s,e:',group_,s,e)
                # print('source:',source)
                source_env=search_dot_2dec(source[i],s,e)  # 列表模式
                try:
                    sim=(e-s)/len(source_env)
                except:
                    sim=0
                sim=round(sim, 3)
                i_,s_,e_=doc2_wrap_dic.get(tem_group)
                target_env = search_dot_2dec(target[i_], s_, e_)

                # for i_ in range(len(doc2_wrap)):
                #     for j_ in range(len(doc2_wrap[i_])):
                #         g_,s_,e_=doc2_wrap[i_][j_]
                #         if g_==tem_group:
                #             target_env = search_dot_2dec(target[i_], s_, e_)#第i段

                source_target_list.append([sim,source_env,target_env])
    source_target_list_sorted = sorted(source_target_list, key=lambda x: x[0], reverse=True)

    result_dup_list=[]
    for i in source_target_list_sorted:
        rate_,source_,target_=i
        if rate_>0 and source_!='' and target_!='' and source_!='\n' and target_!='\n':
            result_dup_list.append({'source':source_,'target':target_,'rate':rate_*100})
    return result_dup_list

def ouput_algo(doc1_wrap):
    s_output_time = time.time()
    # doc1_wrap_2 = []
    # new_old_dic = {}
    # #改写doc1_wrap  <br>融入 前后两组
    # # print('doc1_wrap:', doc1_wrap[0])
    # i = 0
    # # for i,j in enumerate(doc1_wrap[0]):
    # while i < len(doc1_wrap[0]):
    #     # 如果是到最后一个，那就
    #     if (i < len(doc1_wrap[0]) - 2):
    #         a, c, d = doc1_wrap[0][i + 1]  # [(0, 0, 69), (-1, 70, 73), (1, 74, 103)]
    #
    #     if (i < len(doc1_wrap[0]) - 2) and (d - c == 3) and source[a][c:d + 1] == '<br>':  # 前面的一个是<bn>
    #         # #默认中间只有一个分行  如果连续<br>就不好了
    #         a1, c1, d1 = doc1_wrap[0][i]  # 本次
    #         a2, c2, d2 = doc1_wrap[0][i + 2]  # 下下个
    #         doc1_wrap_2.append(tuple([a1, c1, d2]))  # 编号用前面的
    #         new_old_dic[a2] = a1  # a2需要变成a1
    #         i += 3
    #     else:
    #         doc1_wrap_2.append(doc1_wrap[0][i])  # 直接装list
    #         i += 1
    doc1_wrap = [doc1_wrap]
    # # print('合并组之后的doc1_wrap:', doc1_wrap)
    #
    # # print('需要最后改组的:', new_old_dic)
    # # 对wrap2改组号
    # for i, j in enumerate(doc2_wrap[0]):
    #     a, b, c = j
    #     if new_old_dic.get(a, None) != None:
    #         doc2_wrap[0][i] = tuple([new_old_dic[a], b, c])
    # print('改组后的doc2_wrap:', doc2_wrap)
    # print('doc1_wrap是什么?',doc1_wrap) # [(0, 0, 67), (1, 68, 97), (-1, 98, 127), (2, 128, 157)]
    for duan in range(len(doc1_wrap)):
        for num in range(len(doc1_wrap[duan])):
            # print('doc1_wrap[duan][num]是什么?', doc1_wrap[duan][num])
            a, b, c = doc1_wrap[duan][num]
            doc1_wrap[duan][num] = tuple([duan, a, b, c])
    # print('doc1_wrap最后', doc1_wrap[:50])



    print('make output time:', time.time() - s_output_time)
    logging.info('make output time: {}'.format(time.time() - s_output_time))
    return doc1_wrap


def zubao1(x,y,maodian,wrap):
    '''
    x: 是二维的 list
    y: 是一维的str

    '''
    # print('进入组包的wrap',wrap)
    #先做一个阶梯
    yy = [0 for i in range(len(y))]
    n = 0
    for i, j in enumerate(yy):
        # 现在这个位置比n个断电下表要大
        if i >= maodian[n]:
            n += 1
        yy[i] = n * len('<br>')

    for i, j in enumerate(wrap):
        a, b, c = j
        b_p = yy[b]

        c_p = yy[c]
        b += b_p
        c += c_p
        wrap[i] = tuple([a, b, c])

    final_wrap = []
    final_wrap.append(wrap[0])
    for i in range(1, len(wrap)):
        a1, b1, c1 = wrap[i - 1]
        a2, b2, c2 = wrap[i]
        if b2 - c1 > 1:  # 断点  这种断点是因为左右都没有wrap元素
            br_wrap = tuple([-1, c1 + 1, b2 - 1])
            final_wrap.append(br_wrap)
        final_wrap.append(wrap[i])

    join_br='<br>'.join(x)
    # for i, j in enumerate(final_wrap):  # 新的wrap
    #     a, b, c = j
    #     result = yyy[b:c + 1]
    #     print('编号{}，的内容::{}'.format(a, result))

    '''
    final_wrap:[(),()]
    
    '''

    return final_wrap,join_br


def build_label(group_num,content):
    tem = r'<span name="{}">{}</span>'.format(group_num, content) #content是一个字符
    return tem


def zubao2(x,y,maodian,wrap):
    # print('zuboa2:wrap是什么:',wrap)

    '''
    x: 是二维的 list
    y: 是一维的str
    wrap:[] 一维list
    wrap的长度==y

    '''
    # print('进入组包的wrap',wrap)
    #先做一个阶梯
    yy = [0 for i in range(len(y))]
    n = 0
    for i, j in enumerate(yy):
        # 现在这个位置比n个断点下表要大
        if i >= maodian[n]:
            n += 1 # n有可能超出maodian  风险
        yy[i] = n * len('<br>')
    # 上面做好阶梯

    #为每个wrap组调整编号

    global_zihao = []
    for i, j in enumerate(wrap):  # wrap的字号增加
        a, b = j  # 字号，set
        a_p = yy[a]
        a += a_p
        wrap[i] = tuple([a, b])  # 改写了每个wrap中的字号
        global_zihao.append(a)

    #改写完wrap [(字号，set()),()]

    # wrap加入<br>因子
    # 判断连续
    global_y = []
    global_y.append(wrap[0])
    for i in range(1, len(wrap)):
        a1, b1 = wrap[i - 1]
        a2, b2 = wrap[i]
        if a2 - a1 == 5:
            global_y.extend([(a1 + 1, -2), (a1 + 2, -2), (a1 + 3, -2), (a1 + 4, -2)])
            global_y.append(wrap[i])
        elif a2 - a1 == 1:
            global_y.append(wrap[i])
        else:
            print('doc2_wrap重组出错了')
    #wrap加入<br>银子


    # 加入标签
    # 这时候y是不含<br>的，但wrap是有<br>的
    x_br='<br>'.join(x)
    result = ''
    tem = ''
    for i in range(len(wrap)):  #wrap长度应该和y一样
        res1, res2 = wrap[i]
        if res2 == -2:
            result += x_br[i]
        elif res2 == -1:  # -1和-2都是按照原样输入
            result += x_br[i]
        else:  # res是小组
            tem = x_br[i]
            for i in res2:
                if i != -1:
                    tem = build_label(i, tem)
            result += tem
    doc2_str_label=result
    #可以加标签

    # for i, j in enumerate(final_wrap):  # 新的wrap
    #     a, b, c = j
    #     result = yyy[b:c + 1]
    #     print('编号{}，的内容::{}'.format(a, result))

    return doc2_str_label











from my_app.propose.pro1 import gehang

@web.route('/NLP/Algorithm/base/dup_check/winnowing', methods=['POST','GET'])
def dup_check():
    # args_dic = request.args

    global_start_time=time.time()
    print('now route winnowing')
    s_preprocess_time=time.time()
    args_dic=request.form.to_dict()
    print('接收到request:',request)
    print('now time:',time.localtime(time.time()))

    logging.info('foreign request : {} to {}'.format(str(request),'dup_check'))
    key='source'
    sou_key_ok,source,dic=get_key_data(request,key)
    if sou_key_ok==0:
        print("can't get source")
        logging.info("can't get source")
        return jsonify("can't get source")

    key = 'target'
    tem_key_ok, target,dic = get_key_data(request, key)
    if tem_key_ok==0:
        print("can't get target")
        logging.info("can't get target")
        return jsonify("can't get target")


    template_target = dic.get('template','') #有template 或者没有
    # print('提取的模板:',template_target)
    template_length=len(template_target)
    # source , target template 的异常[]   [……[]]
    # str阶段 ''或者无，'……'

    #分段并且去掉空段
    # template_target = template_target.split(r'\n')
    tem_str=template_target
    template_target=[template_target] # '' 变成[''] #含有换行符


    template_target = clear(template_target) # 其实clear没用  看看需不需要 my_split
    # print('clear之后的template_target:',template_target)

    source_length = len(source)
    target_length = len(target)
    print('长度：  source: {} | target : {} | template: {}'.format(source_length,target_length,template_length))

    # print('clear后的内容:')
    # print('source:',source[:100] )
    # print('target:', target[:100])
    # print('tem:',template_target[:100])
    print('tem_string:',tem_str[:100])
    print('tem_str是什么?123',repr(tem_str[:100])) # 应该是None

    tem_fenduan, tem_split, tem_duandian = my_split(tem_str) # 输入是str  先分段，然后去掉空行 然后返回拼接或者直接返回段信息
    # x_fenduan 是一维的，每维是一段一个str
    # print('tem_fenduan是什么?',tem_fenduan[:5]) # ['']

    # print('source原文：',repr(source))
    sour_gechang_time=time.time()
    source=gehang(source)
    print('source gehang时间:',time.time()-sour_gechang_time)

    sou_ms=time.time()
    x_fenduan,source,x_duandian=my_split(source) # str
    print('sou my split time:',time.time()-sou_ms)

    # print('去掉泛函之后的文本:',source)

    # print('字符串source:',repr(source[:1000]))

    target = gehang(target)
    y_fenduan,target,y_duandian=my_split(target) # str  <br>连起来
    # print('字符串target:', repr(target[:1000]))
    # print('去掉泛函之后的文本target:', target)
    # source = source.replace('\n', '<br>')
    # print('replace之后source', source)
    # target = target.replace('\n', '<br>')
    source=[source] #
    target=[target]
    # source=clear(source)#去掉空段之后，至少存在一个['']
    # target = clear(target)
    # print('clear之后的source:',source[:1000])
    # print('clear之后的source:', source[:1000])

    print('preprocess time:',time.time()-s_preprocess_time)
    logging.info('preprocess time: {}'.format(time.time()-s_preprocess_time))
    s_time=time.time()
    example = paragraph_winnowing()
    # print('送入检测的source:',len(source[0]),source)  #103长度
    # print('送入检测的target:', target)

    # print('连续的source:',source[0][:5000])
    # print('连续的target:', target[0])
    # sen_exam=r'我们完全理解并同意放弃对这方面有不明及误解的权利'
    # my_leng=len(sen_exam)
    # posi_exam_s=[]
    # posi_exam_t=[]
    # for i in range(len(source)):
    #     if len(source[i:i+my_leng])!=my_leng:
    #         print('长度不对')
    #     if source[i:i+my_leng]==sen_exam:
    #         print('source找到所在位置',i)
    #         posi_exam_s.append(i)
    #
    # for i in range(len(target)):
    #     if target[i:i+my_leng]==sen_exam:
    #         print('target找到所在位置',i)
    #         posi_exam_t.append(i)
    # if posi_exam_s:
    #     print('juzisuozai huanjing_s:',source[posi_exam_s[0]-50:posi_exam_s[0]+100])
    # if posi_exam_t:
    #     print('juzisuozai huanjing_t:', target[posi_exam_t[0] - 50:posi_exam_t[0] + 100])




    similarity,result_str,doc1_wrap,doc2_wrap=example.get_sim(source,target,template=template_target,n=13)

    # print('getsim的doc2_wrap:',doc2_wrap)

    # 第二项其实没用到
    # print('未加入br的wrap1:', doc1_wrap) # 下表最大是102  [[(0, 0, 101), (-1, 102, 102)]]

    # result_dup_list = list_model(doc1_wrap, doc2_wrap, source, target)
    # print('result_dup_list:',result_dup_list)


    # [{'source': '我是马大哈我是马大哈我是马大哈我是马大哈我是马大哈我是马大哈我是马大哈我是马大哈。我是梁静怡我是梁静怡我是梁静怡我是梁静怡我是梁静怡我是梁静怡。', 'target': '我是梁静怡我是梁静怡我是梁静怡我是梁静怡我是梁静怡我是梁静怡哈哈哈哈哈。', 'rate': 98.6}]

    # 给source和target加入br
    x_final_wrap,x_join_br = zubao1(x_fenduan,source[0],x_duandian,doc1_wrap[0])
    # print('组包之后x_final_wrap',x_final_wrap)
    # print('x_join_br:',x_join_br)

    # 给target加入br
    doc2_str_label = zubao2(y_fenduan, target[0], y_duandian, doc2_wrap)
    # 直接输出就行


    # print('x_final_wrap是什么?',x_final_wrap)
    # print('x_join_br是什么?', x_join_br)
    # print('y_final_wrap是什么?', y_final_wrap)
    # print('y_join_br是什么?', y_join_br)

    # source_dup_dict=source_dup_dic(result_str)
    print('get sim run time :',time.time()-s_time)
    print('similarity:', similarity)

    # print('ouput_algo之前的doc1_wrap:',doc1_wrap[:20])
    # print('x_join_br:',x_join_br)
    # print('抽出来看:',x_join_br[0][782:800])

    x_final_wrap=ouput_algo(x_final_wrap)  #把wrap中的 <br>一下，正常来说，通过zubao，是不用改的

    # print('output之后的x_final_wrap',x_final_wrap) # [[(0, 0, 0, 67), (0, 1, 68, 97), (0, -1, 98, 127), (0, 2, 128, 157)]]
    # print('x_join_br:',x_join_br)
    x_join_br=[x_join_br]



    # print('ouput_algo之后的doc1_wrap:', x_final_wrap[:20])


    # for i,j in enumerate(source_target_list_sorted):#加上编号
    #     source_target_list_sorted[i].insert(0,i)
    # print('加入编号后的list',source_target_list_sorted)
    # print('make output time:',time.time()-s_time)


    # make source target 字典
    # source_dic={}
    # target_dic={}
    # group__=0
    # for i in range(len(source_target_list)):
    #     source_tem,target_tem=source_target_list[i]
    #     source_dic[group__]=source_tem
    #     target_dic[group__] = target_tem
    #     group__+=1
    #
    # print('source_dic::::',source_dic)
    # print('target_dic::::', target_dic)

    # 隔段合并的问题:

    time_=time.time()-global_start_time
    print('全局时间:',time_)
    logging.info('run success!! time cost      :    {}      |length : {}  |  {}  |    {}  |dup_rate:{}'.format(time_,
                                                                                                               source_length,
                                                                                                               target_length,
                                                                                                               template_length,
                                                                                                               similarity))

    result1=similarity

    result3 = render_template('testHtml2.html', name1='doc1', name2='doc2', time=time_, dup_check=similarity,doc1_str=x_join_br, doc2_str=doc2_str_label,
                    doc1_wrap=x_final_wrap)
    # return result3

    # 左边文本
    result4 = render_template('add_href_doc1.html', doc1_wrap=x_final_wrap, doc1_str=x_join_br)


    # # return result3
    #右边文本

    print('doc2_str_label:',doc2_str_label[:5000])

    result5 = render_template('add_href_doc2.html',doc2_str=doc2_str_label)


    # # result6 = render_template('dup_list_source.html', source_dup=source_target_list_sorted)
    # # result7 = render_template('dup_list_target.html', target_dup=source_target_list_sorted)
    #
    result_dic = {'dup_rate': result1,
                  'source_label': result4,
                  'target_label': result5
                  }
    #
    return jsonify(result_dic)



    #
    #
    # time_, result = check_str(str(doc1), str(doc2), k=13)
    # similiarize, dup_text, dup_dic, doc1_str, doc2_str, doc1_wrap, doc2_wrap = result
    #
    # result1 = str(similiarize)
    # result2 = dup_dic
    # result3 = render_template('testHtml.html', name1='doc1', name2='doc2', time=time_, dup_check=similiarize,
    #                           dup_text=dup_text, dup_dic=dup_dic, doc1_str=doc1_str, doc2_str=doc2_str,
    #                           doc1_wrap=doc1_wrap, doc2_group_=doc2_wrap)
    # result4 = render_template('add_href_doc1.html', doc1_wrap=doc1_wrap, doc1_str=doc1_str)
    # result5 = render_template('add_href_doc2.html', doc2_group_=doc2_wrap, doc2_str=doc2_str)
    #
    # print('重复文本字典：',result2)
    # result6 = time_
    # result7 = doc1_str
    # result = doc2_str
    # # return result4
    #
    # return result4
    #
    # result_dic = {'dup_rate': result1,
    #               'doc1_label': result4,
    #               'doc2_label': result5}
    # return jsonify(result_dic)


def get_doc(request,key,content_ok=0):
    try:
        try:
            dic = request.args.to_dict()
            content = dic[key]
            content_ok = 1
            content = json.loads(content)
            # print('{}内容为:{}'.format(key,content))
        except:
            dic = request.form.to_dict()  #
            content = dic[key]
            content_ok = 1
            content = json.loads(content)
            # print('{}内容为:{}'.format(key,content))
    except:
        content=None
    return content_ok,content

def intepret_docx(url):
    res = requests.post(url)
    Byio = BytesIO(res.content)

    content = docx.Document(Byio)
    return content


def propose_docx_doc(source_url,template_url):
    source_doc_name=source_url.split('/')[-1]
    tem_doc_name = template_url.split('/')[-1]
    source_isdoc = 0
    tem_isdoc = 0
    if source_doc_name.endswith('doc'): #是doc文件
        # source_content_ok = 0
        source_isdoc=1
        key='source_content'
        source_content_ok,source_content=get_doc(request,key) #source_content 正确或者None

    else:# 是docx文件
        source_content=intepret_docx(source_url)
        # '招标文件-基于NLP的商务文本数据清洗关键技术研究（1021评审）（第一版）.docx'
        # path1 = r"D:\lcb_note\code\Program\10月项目\查重需求资料\查重需求资料\1.招标文件文档查重对比材料\招标文件 CWEME-1911ZSWZ-2J039 基于NLP的商务文本数据清洗关键技术研究项目-2019年12月中国水利电力物资集团有限公司项目（第三版终版）.docx"
        # path1=r'D:\lcb_note\code\Program\10月项目\winnowing1.5\越界1\招标文件 CWEME-1910ZSWZ-2J036 物资成套信息管理平台招标-2019年10月中国水利电力物资集团有限公司项目-招标三部（第二版模板）.doc'
        # path1=r'D:\lcb_note\code\Program\10月项目\winnowing1.5\越界1\招标文件-基于NLP的商务文本数据清洗关键技术研究（1021评审）（第一版）.docx'
        # path1=r'D:\lcb_note\code\Program\10月项目\my_docx\基于NLP的商务文本数据清洗关键技术研究项目合同+-+-打印版.docx'
        # source_content = docx.Document(path1)

    if tem_doc_name.endswith('doc'):
        tem_isdoc=1
        # template_content_ok = 0
        key='template_content'
        tem_content_ok, template_content = get_doc(request, key)
        # print('template_content是什么?',template_content)
    else:# docx文件
        template_content = intepret_docx(template_url)
        # path2 = r'D:\lcb_note\code\Program\10月项目\查重需求资料\查重需求资料\1.招标文件文档查重对比材料\招标文件-基于NLP的商务文本数据清洗关键技术研究（1021评审）（第一版）.docx'
        # path2=r'D:\lcb_note\code\Program\10月项目\winnowing1.5\越界1\招标文件 CWEME-1910ZSWZ-2J036 物资成套信息管理平台招标-2019年10月中国水利电力物资集团有限公司项目-招标三部（第二版模板）.doc'
        # path2 = r'D:\lcb_note\code\Program\10月项目\my_docx\招标文件 CWEME-1911ZSWZ-2J039 基于NLP的商务文本数据清洗关键技术研究项目-2019年12月中国水利电力物资集团有限公司项目（第三版终版）.docx'
        # template_content = docx.Document(path2)

    return source_content,template_content,source_doc_name,tem_doc_name,source_isdoc,tem_isdoc



@web.route('/NLP/Algorithm/base/dup_check/simhash', methods=['POST','GET'])
def simhash_route():
    # args_dic = request.args
    global_start_time=time.time()

    print('########################  now route simhash  ###############################')
    # s_preprocess_time=time.time()
    # args_dic=request.form.to_dict()
    print('接收到request:',request)
    print('now time:',time.localtime(time.time()))
    logging.info('{} foreign request : {} to {}'.format('simhash:',str(request),'dup_check'))
    key='source'
    sou_key_ok,source,dic=get_key_data(request,key)
    if sou_key_ok==0:
        print("can't get source")
        logging.info("can't get source")
        return jsonify("can't get source")
    key = 'target'
    tem_key_ok, target, dic = get_key_data(request, key)
    if tem_key_ok == 0:
        print("can't get target")
        logging.info("can't get target")
        return jsonify("can't get target")

    template_target = dic.get('template', '')  # 有template 或者没有
    # print('提取的模板:',template_target)
    template_length = len(template_target)
    # source , target template 的异常[]   [……[]]
    # str阶段 ''或者无，'……'

    # 分段并且去掉空段
    # template_target = template_target.split(r'\n')
    tem_str = template_target # '含符号'
    print('模板加入:',tem_str[:100])
    # template_target = [template_target]
    # template_target = clear(template_target)  # 其实clear没用
    # print('clear之后的template_target:',template_target)

    source_length = len(source)
    target_length = len(target)
    print('长度：  source: {} | target : {} | template: {}'.format(source_length, target_length, template_length))

    tem_fenduan, tem_split, tem_duandian = my_split(tem_str)
    x_fenduan, source, x_duandian = my_split(source)
    y_fenduan, target, y_duandian = my_split(target)

    # simhash
    simh_time_s = time.time()
    sim_list = sim_main(x_fenduan, y_fenduan, tem_fenduan)
    print('simhash全部时间:', time.time() - simh_time_s)
    dup_list_simhash = []
    for i, j in enumerate(sim_list):
        rate, doc1_index, dis, doc2_index, doc1, doc2 = j
        dic_sim = {'source': doc1, 'target': doc2, 'rate': rate}
        dup_list_simhash.append(dic_sim)
    # print('组装好的dup_list_simhash:',dup_list_simhash[:10])

    result_dup_list = dup_list_simhash
    result_dic={'dup_list':result_dup_list}
    return jsonify(result_dic)


from my_app.algorithm.template_match.algo3 import main
@web.route('/NLP/Algorithm/base/dup_check/template_match', methods=['POST','GET'])
def template_match():
    #接收区
    logging.info('foreign request : {} to {}'.format(str(request),'template_match'))
    print('接收到request:', request)
    print('now time:', time.localtime(time.time()))
    # print('Content-Type:',request.content_encoding)# dir(request),
    source_ok,template_ok,source_url,template_url=vali_check_match1(request)

    if source_ok == 0:
        print("can't get source")
        logging.info("can't get source")
        return jsonify("can't get source")
    if template_ok == 0:
        print("can't get template")
        logging.info("can't get template")
        return jsonify("can't get template")


    source_content,template_content,source_doc_name,tem_doc_name,source_isdoc,tem_isdoc=propose_docx_doc(source_url,template_url)

    print('两个文档:', source_doc_name, tem_doc_name)
    print('各种状态值:source_isdoc:{},tem_isdoc:{}'.format(source_isdoc,tem_isdoc))

    if not source_content:return jsonify("can't get source_content")
    if not template_content: return jsonify("can't get template_content")

    print('source_url是什么?',source_url)
    print('template_url是什么?', template_url)

    start_time=time.time()
    left,right,tem_global_list_obj,source_global_obj_list,match_rate_head=main(source_content,template_content,source_isdoc,tem_isdoc)

    # print('tem_global_list_obj:',tem_global_list_obj)
    # print('source_global_obj_list:',source_global_obj_list)

    # left_=[dict(i._asdict()) for i in left]
    # right_ = [dict(i._asdict()) for i in right]
    # print('left_:',left_)
    # print('right:',right_)

    # print('tem_global_list_obj:',tem_global_list_obj)
    # print('source_global_obj_list:',source_global_obj_list)
    result0=render_template('html_for_match.html',template=tem_global_list_obj )
    result1 = render_template('html_for_match_right.html', template=source_global_obj_list)

    result_dic={
        'match_rate':match_rate_head,
        'template':result0,
        'source':result1
    }

    # return result1
    print('整个过程时间:',time.time()-start_time)
    return jsonify(result_dic)  #obj传不了

# 'source_info': '1(正确),-2(多余),-3(位置不正确),-4(位置正确但级别不对)',
# 'tem_info': '1(正确),-2(缺失),-3(位置不正确),-4(位置正确但级别不对)',

# <p style="color:red">红色的字</p>


# def func(x):
#     print('x是什么?',x)
#     print('我是{}'.format(x))

# from my_app.web.test2 import parallelize
# @web.route('/NLP/Algorithm/base/dup_check/duojincheng', methods=['POST','GET'])
# def duojin():
#     # key='source'
#     # key_ok,content,dic=get_key_data(request,key)
#     data_input = [[[1], [2]], [[3], [4]]]  # 可以进入func再解包   每次读入 [[1],[2]]
#     parallelize(data_input, func)
#     return 'ok'













































