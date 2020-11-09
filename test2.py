# clear之后的source: [' ', '基于 NLP 的商务文本数据清洗关键技术', '研究项目招标-2019 年 12 月中国水利电', '力物资集团有限公司项目 ', ' ', '招标编号：', 'CWEME-1912ZSWZ-2J039(二次) ', ' ', '技术文件 ', ' ', '\t', ' ', '投标人：北京科东电力控制系统有限责任公司（盖单位章） ', '法定代表人或其委托代理人：          （签字） ', '2019 年 12 月 10 日 ', ' ', ' ', '\t', ' ', '\t 2\t', '目录 ', '一.对招标范围及要求的理解\t.........................................................................................................\t5\t', '1. 工作目 标\t............................................................................................................................\t5\t', '2. 招标范围及内容\t.................................................................................................................\t5\t', '2.1 文本数据挖掘关键理论与技术研究\t........................................................................\t5\t', '2.2 商务文本数据清洗原型系统开发\t............................................................................\t6\t', '2.3 商务文本数据挖掘原 型系统实证测试\t....................................................................\t6\t', '2.4 研究成果交付\t............................................................................................................\t6\t', '3. 工作要求\t............................................................................................................................\t6\t', '4. 投标技术文件内容要求\t.....................................................................................................\t7\t', '二.对项目研究技术规范的理解\t.....................................................................................................\t8\t', '1. 企业概况\t............................................................................................................................\t8\t', '2. 项目研究目标的实现与应用实践\t.....................................................................................\t8\t', '2.1 研究原则\t....................................................................................................................\t8\t', '2.2 研究目标\t....................................................................................................................\t8\t', '2.3 项目范围\t....................................................................................................................\t9\t', '2.4 交付要求和范围\t........................................................................................................\t9\t', '2.5 项目研究计划及要求\t..............................................................................................\t10\t', '2.6 投标的基本技术要求\t..............................................................................................\t10\t', '3. 原型系统性能要求\t...........................................................................................................\t13\t', '4. 平台技术要求\t...................................................................................................................\t13\t', '4.1 基本要求\t..................................................................................................................\t13\t', '4.2 原型系统架构及数据库要求\t..................................................................................\t14\t', '4.3 原型系统设计要求\t..................................................................................................\t14\t', '4.4 系统维护要求\t..........................................................................................................\t15\t', '5. 课题关键技术攻关要求\t...................................................................................................\t15\t', '6. 培训要求\t..........................................................................................................................\t15\t', '6.1 总体要求\t..................................................................................................................\t15\t', '6.2 培训对象和要求\t......................................................................................................\t16\t', '7. 项目验收\t..........................................................................................................................\t16\t']
import docx
import win32
import win32com
# from pywin32.win32com.client import Dispatch
path1=r'D:\lcb_note\code\Program\10月项目\winnowing1.5\pri_lcb\123.doc'
path2=r'.\招标文件 CWEME-1911ZSWZ-2J039 基于NLP的商务文本数据清洗关键技术研究项目-2019年12月中国水利电力物资集团有限公司项目（第三版终版）.docx'

# word=

docx.Document(path2)

# def doc2docx(path):
#     w = win32com.client.Dispatch('Word.Application')
#     w.Visible = 0
#     w.DisplayAlerts = 0
#     doc = w.Documents.Open(path)
#     newpath = os.path.splitext(path)[0] + '.docx'
#     doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
#     doc.Close()
#     w.Quit()
#     os.remove(path)
#     return newpath
import subprocess
import os
import docx

filename=path1
if filename.endswith('.doc'):

    subprocess.call(['soffice', '--headless', '--convert-to', 'docx', filename])

    doc = docx.Document(filename[:-4]+".docx")
    for para in doc.paragraphs:
        print (para.text)



# doc1=doc2docx(path1)
# docx.Document(doc1)











